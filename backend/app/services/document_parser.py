"""
Document parsing for Agent 1's "Contextual Knowledge" upload tab.
Extracts plain text from .docx, .pdf, .txt, .md files for chunking + embedding.
"""
import logging
import io

logger = logging.getLogger(__name__)


def extract_text(filename: str, content: bytes) -> str:
    """Extract plain text from uploaded file bytes based on extension."""
    name = (filename or "").lower()

    if name.endswith(".docx"):
        return _extract_docx(content)
    if name.endswith(".pdf"):
        return _extract_pdf(content)
    if name.endswith((".txt", ".md")):
        try:
            return content.decode("utf-8", errors="ignore")
        except Exception:
            return content.decode("latin-1", errors="ignore")

    # Unknown type — try plain decode as a last resort
    try:
        return content.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"docx extraction error: {e}")
        return ""


def _extract_pdf(content: bytes) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))
        parts = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text.strip())
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"pdf extraction error: {e}")
        return ""
